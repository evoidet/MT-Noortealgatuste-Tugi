#!/usr/bin/env python3
"""Prepare private DOCX templates from the approved reference documents.

The source files are read-only design authorities. This script writes new,
clean server-side templates: sample covers, example markers, filling guidance,
placeholder prose, and the internal decision page are intentionally removed.
"""

from __future__ import annotations

import argparse
import hashlib
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


DEFAULT_INVOICE_SOURCE = Path.home() / "Downloads" / "MTU_Noortealgatuste_Tugi_arve_naidis.docx"
DEFAULT_EXPENSE_SOURCE = Path.home() / "Downloads" / "Naidisdokument_kulude_huvitamise_avaldus_ja_kuluaruanne.docx"
DEFAULT_OUTPUT_ROOT = Path(__file__).resolve().parents[1] / "private" / "templates" / "documents"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_paragraph(paragraph) -> None:
    remove_element(paragraph._p)


def remove_table(table) -> None:
    remove_element(table._tbl)


def remove_row(table, index: int) -> None:
    remove_element(table.rows[index]._tr)


def set_run_plain(run, *, bold: bool | None = None, color: str = "111111") -> None:
    run.bold = bold
    run.italic = False
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_text(paragraph, text: str, *, bold: bool | None = None, color: str = "111111") -> None:
    """Replace visible text while preserving paragraph and first-run geometry."""
    runs = paragraph.runs
    if not runs:
        run = paragraph.add_run()
    else:
        run = runs[0]
        for extra in runs[1:]:
            remove_element(extra._r)
    run.text = text
    set_run_plain(run, bold=bold, color=color)


def set_paragraph_runs(paragraph, parts: list[tuple[str, bool]]) -> None:
    for run in list(paragraph.runs):
        remove_element(run._r)
    for text, bold in parts:
        run = paragraph.add_run(text)
        set_run_plain(run, bold=bold)


def set_cell_tag(cell, tag: str, *, bold: bool | None = None, color: str = "111111") -> None:
    set_paragraph_text(cell.paragraphs[0], tag, bold=bold, color=color)


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def prepare_invoice(source: Path, destination: Path) -> None:
    document = Document(source)

    # The first section is a sample cover containing only a giant "ARVE".
    # Its section properties live on the first body paragraph; removing that
    # paragraph makes the real invoice page the only section/page.
    remove_paragraph(document.paragraphs[0])

    header, metadata, parties, items, totals, payment = document.tables

    metadata_tags = {
        (0, 1): "{invoiceNumber}",
        (0, 3): "{invoiceDate}",
        (1, 1): "{dueDate}",
        (1, 3): "{currency}",
        (2, 1): "{transactionTime}",
        (2, 3): "{projectReference}",
    }
    for (row, column), tag in metadata_tags.items():
        set_cell_tag(metadata.cell(row, column), tag)

    buyer = parties.cell(0, 1)
    set_paragraph_text(buyer.paragraphs[1], "{buyerName}", bold=True)
    set_paragraph_runs(buyer.paragraphs[2], [("Registrikood: ", False), ("{buyerRegistryCode}", False)])
    set_paragraph_runs(buyer.paragraphs[3], [("Aadress: ", False), ("{buyerAddress}", False)])
    set_paragraph_runs(buyer.paragraphs[4], [("Kontaktisik: ", False), ("{buyerContact}", False)])

    # One tagged row is duplicated by docxtemplater for any number of line items.
    for index in range(len(items.rows) - 1, 1, -1):
        remove_row(items, index)
    row = items.rows[1]
    item_tags = [
        "{#items}{number}",
        "{description}",
        "{quantity}",
        "{unit}",
        "{unitPrice}",
        "{lineTotal}{/items}",
    ]
    for cell, tag in zip(row.cells, item_tags, strict=True):
        set_cell_tag(cell, tag)

    set_cell_tag(totals.cell(0, 1), "{subtotal}", bold=True)
    set_cell_tag(totals.cell(1, 1), "{vatText}")
    set_cell_tag(totals.cell(2, 1), "{total}", bold=True, color="FFFFFF")

    payment_cell = payment.cell(0, 0)
    set_paragraph_runs(payment_cell.paragraphs[3], [("Selgitus: ", False), ("{paymentDescription}", False)])
    set_paragraph_runs(payment_cell.paragraphs[4], [("Viitenumber: ", False), ("{referenceNumber}", False)])

    for paragraph in list(document.paragraphs):
        if "T\u00c4ITMISE ABI" in paragraph.text.upper() or "KUSTUTA ENNE SAATMIST" in paragraph.text.upper():
            remove_paragraph(paragraph)

    set_update_fields(document)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def prepare_expense(source: Path, destination: Path) -> None:
    document = Document(source)

    guidance = document.tables[0]
    general = document.tables[1]
    activity = document.tables[2]
    costs = document.tables[3]
    signature = document.tables[4]

    # Everything from the page break preceding the internal decision onward
    # is outside the ordinary applicant-facing report.
    page_break_before_decision = document.paragraphs[24]
    body = document._element.body
    children = list(body.iterchildren())
    start_index = children.index(page_break_before_decision._p)
    for child in children[start_index:]:
        if child.tag != qn("w:sectPr"):
            remove_element(child)

    remove_table(guidance)

    general_tags = {
        0: "{documentNumberAndDate}",
        4: "{recipientName}",
        5: "{recipientRole}",
        6: "{contactAccountIban}",
        7: "{activityName}",
        8: "{expenseType}",
        9: "{locationPeriodRoute}",
        10: "{fundingSource}",
    }
    for row, tag in general_tags.items():
        set_cell_tag(general.cell(row, 1), tag)

    activity_tags = ["{whereWhen}", "{activitiesAndRole}", "{necessity}", "{result}", "{participants}"]
    for row, tag in enumerate(activity_tags):
        set_cell_tag(activity.cell(row, 1), tag)

    # Remove the filling-rules box and spare example rows, retaining one loop
    # row plus the totals row.
    remove_row(costs, 7)
    for index in range(5, 1, -1):
        remove_row(costs, index)
    cost_row = costs.rows[1]
    cost_tags = [
        "{#items}{description}",
        "{date}",
        "{documentReference}",
        "{grossAmount}",
        "{requestedAmount}",
        "{excludedAmount}{/items}",
    ]
    for cell, tag in zip(cost_row.cells, cost_tags, strict=True):
        set_cell_tag(cell, tag)
    total_row = costs.rows[2]
    for column, tag in [(3, "{grossTotal}"), (4, "{requestedTotal}"), (5, "{excludedTotal}")]:
        set_cell_tag(total_row.cells[column], tag)

    # Replace the example fragments in the application sentence with fields.
    application = document.paragraphs[8]
    set_paragraph_runs(
        application,
        [
            (
                "Palun h\u00fcvitada mulle eespool nimetatud MT\u00dc p\u00f5hikirjalise tegevusega seotud ja "
                "dokumentaalselt t\u00f5endatud kulud kokku ",
                False,
            ),
            ("{requestedTotal}", True),
            (" arvelduskontole ", False),
            ("{iban}", True),
            (".", False),
        ],
    )

    set_cell_tag(signature.cell(1, 0), "{recipientName}")
    set_cell_tag(signature.cell(1, 1), "{signatureStatus}")
    set_cell_tag(signature.cell(1, 2), "{signatureDate}")

    # Replace sample attachment guidance with a real, data-driven list.
    set_paragraph_text(document.paragraphs[17], "Lisatud dokumendid:")
    set_paragraph_text(document.paragraphs[18], "{#attachments}")
    set_paragraph_text(document.paragraphs[19], "{name}")
    set_paragraph_text(document.paragraphs[20], "{/attachments}")
    for paragraph in list(document.paragraphs[21:24]):
        remove_paragraph(paragraph)

    # Header/footer sample markers and instructions are not part of the form.
    header_table = document.sections[0].header.tables[0]
    set_cell_tag(header_table.cell(0, 0), "")
    footer_table = document.sections[0].footer.tables[0]
    set_cell_tag(footer_table.cell(0, 0), "")

    set_update_fields(document)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--invoice-source", type=Path, default=DEFAULT_INVOICE_SOURCE)
    parser.add_argument("--expense-source", type=Path, default=DEFAULT_EXPENSE_SOURCE)
    parser.add_argument("--output-root", type=Path, default=DEFAULT_OUTPUT_ROOT)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    invoice_source = args.invoice_source.resolve(strict=True)
    expense_source = args.expense_source.resolve(strict=True)
    invoice_before = sha256(invoice_source)
    expense_before = sha256(expense_source)

    invoice_destination = args.output_root / "arve" / "arve.docx"
    expense_destination = args.output_root / "kuluaruanne" / "kuluaruanne.docx"
    prepare_invoice(invoice_source, invoice_destination)
    prepare_expense(expense_source, expense_destination)

    if sha256(invoice_source) != invoice_before or sha256(expense_source) != expense_before:
        raise RuntimeError("A reference DOCX changed while preparing templates")

    print(f"Invoice template: {invoice_destination.resolve()}")
    print(f"Expense template: {expense_destination.resolve()}")
    print(f"Invoice source SHA-256: {invoice_before}")
    print(f"Expense source SHA-256: {expense_before}")


if __name__ == "__main__":
    main()
